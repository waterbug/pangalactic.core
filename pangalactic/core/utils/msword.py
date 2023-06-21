# -*- coding: utf-8 -*-
"""
Pan Galactic report writer to produce Word documents.
"""
import os
from docx import Document
from docx.shared import Inches

from pangalactic.core.uberorb import orb


def report(headers, data, info=None, output_path=''):
    template_path = os.path.join(orb.vault, 'landscape.docx')
    doc = Document(template_path)
    if info:
        doc.add_heading(f'CATTENS Hardware Library: {info}', 0)
    else:
        doc.add_heading('CATTENS Hardware Library', 0)
        info = "all"
    table = doc.add_table(rows=1, cols=len(headers),
                          style='Light List Accent 1')
    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
    for i, row in enumerate(data):
        row_cells = table.add_row().cells
        for j, val in enumerate(row):
            row_cells[j].text = str(val)
    fpath = os.path.join(output_path, f'{info}.docx')
    doc.save(fpath)


def demo():
    # output a demo doc ...
    document = Document()
    document.add_heading('Document Title', 0)
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    document.add_paragraph('first item in unordered list', style='List Bullet')
    document.add_paragraph('first item in ordered list', style='List Number')
    document.add_picture('monty-truth.png', width=Inches(1.25))
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    table = document.add_table(rows=1, cols=3, style='Light List Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    document.add_page_break()
    document.save('/home/waterbug/demo.docx')


if __name__ == '__main__':
    demo()

